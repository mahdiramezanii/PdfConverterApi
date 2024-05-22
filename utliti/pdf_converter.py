from django.utils.crypto import get_random_string
from tqdm import tqdm
from pdf2image import convert_from_path
import pytesseract
from PIL import Image
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from pathlib import Path
import tempfile
import sys


def pdf_to_word(pdf_path:str, output_dir:str, lang='fas', **kwargs):
    """
    convert pdf to word
    """

    output_dir=output_dir.replace("\\","/")
    print(output_dir)
    # pdf_path = Path(pdf_path)
    pdf_name = f"word-{get_random_string(5)}"
    pages = convert_from_path(pdf_path,poppler_path="C:/Program Files (x86)/poppler-24.02/bin")

    pytesseract.pytesseract.tesseract_cmd = "C:/Program Files/Tesseract-OCR/tesseract.exe"

    texts = []

    for i, page in tqdm(enumerate(pages), position=0):

        with tempfile.TemporaryDirectory() as img_dir:
            img_name = f'{pdf_name}-{i+1}.jpg'
            img_path = Path(img_dir) / img_name

            page.save(img_path, 'JPEG')
            text = pytesseract.image_to_string(Image.open(img_path), lang="fas+eng")
            texts.append(text)

    document = Document()
    style_normal = document.styles['Normal']
    font = style_normal.font
    font.name = 'Arial'
    font.rtl = True

    style_h1 = document.styles['Heading 1']
    font = style_h1.font
    font.name = 'Arial'
    font.rtl = True

    for i, text in tqdm(enumerate(texts), position=0):
        heading = document.add_heading(f'صفحه: {i+1}', level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        heading.style = document.styles['Heading 1']

        paragraph = document.add_paragraph(text)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        paragraph.style = document.styles['Normal']

    output_path = Path(output_dir) / f'{pdf_name}.docx'
    document.save(output_path)

    return f'{pdf_name}.docx'


def convert_url(file_path:str,directory_path:str):


    directory_path=directory_path.replace("\\","/")

    test=directory_path.split("/")
    directory_path_edited=""

    for index in range (len(test)-1):
         if test[index] != "media":
            directory_path_edited+=test[index]
         if ((index+1) < (len(test))-1):
             directory_path_edited += "/"


    result=directory_path_edited+file_path

    return result



